"""Extract plain text from an uploaded project specification.

Accepts PDF, Word (.doc/.docx), Markdown, or plain text. Used so a researcher
can drop in a grant aim, proposal, or one-paragraph description and have each
paper assessed against it.
"""
from __future__ import annotations

import io
import os
import re
import shutil
import subprocess
import tempfile


def extract_text(filename: str, raw: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _from_pdf(raw)
    if name.endswith(".docx"):
        return _from_docx(raw)
    if name.endswith(".doc"):
        return _from_doc(raw)
    # Markdown / txt / anything else: decode as UTF-8, tolerant of junk bytes.
    return raw.decode("utf-8", errors="replace").strip()


def _from_pdf(raw: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("pypdf is required to read PDF specs.") from exc

    reader = PdfReader(io.BytesIO(raw))
    chunks = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception:
            continue
    text = "\n\n".join(chunks).strip()
    if not text:
        raise RuntimeError("Could not extract any text from this PDF (it may be scanned images).")
    return text


def _from_docx(raw: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required to read Word specs.") from exc

    doc = Document(io.BytesIO(raw))
    parts: list[str] = []
    for para in doc.paragraphs:
        line = para.text.strip()
        if line:
            parts.append(line)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n\n".join(parts).strip()
    if not text:
        raise RuntimeError("Could not extract any text from this Word document.")
    return text


def _from_doc(raw: bytes) -> str:
    """Legacy binary .doc — try antiword/catdoc, then OLE text recovery."""
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(raw)
        path = tmp.name
    try:
        for cmd in (["antiword", path], ["catdoc", path]):
            if shutil.which(cmd[0]):
                try:
                    result = subprocess.run(
                        cmd,
                        capture_output=True,
                        text=True,
                        timeout=30,
                        check=False,
                    )
                except (OSError, subprocess.TimeoutExpired):
                    continue
                text = (result.stdout or "").strip()
                if result.returncode == 0 and text:
                    return text
        return _from_doc_ole(raw)
    finally:
        try:
            os.unlink(path)
        except OSError:
            pass


def _from_doc_ole(raw: bytes) -> str:
    try:
        import olefile
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("olefile is required to read legacy .doc specs.") from exc

    if not olefile.isOleFile(io.BytesIO(raw)):
        raise RuntimeError("Not a valid .doc file.")
    ole = olefile.OleFileIO(io.BytesIO(raw))
    if not ole.exists("WordDocument"):
        raise RuntimeError("Could not read this .doc file. Try saving as .docx or paste the text.")
    data = ole.openstream("WordDocument").read()
    pieces: list[str] = []
    for match in re.finditer(rb"(?:[\x20-\x7e\r\n\t]\x00){4,}", data):
        try:
            pieces.append(match.group(0).decode("utf-16le"))
        except Exception:
            continue
    text = "\n".join(pieces).strip()
    if len(text) < 20:
        raise RuntimeError(
            "Could not extract text from this .doc file. Save as .docx or paste the text."
        )
    return text
